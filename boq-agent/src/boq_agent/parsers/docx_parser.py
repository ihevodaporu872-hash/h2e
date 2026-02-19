"""Word document (.docx) parser."""

from pathlib import Path
from typing import Optional
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table

from .base import (
    BaseParser,
    DocumentContent,
    ParsedTable,
    ParserErrorType,
    ParserResult,
)


class DocxParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: Path) -> ParserResult:
        """Parse Word document and extract content."""
        if not file_path.exists():
            return self._create_error_result(
                file_path,
                ParserErrorType.FILE_NOT_FOUND,
                f"File not found: {file_path}",
            )

        result = ParserResult(
            file_path=file_path,
            file_type=".docx",
            success=True,
        )

        try:
            doc = Document(file_path)
            self._extract_metadata(doc, result)
            self._extract_content(doc, result)
            self._extract_tables(doc, result)

        except PackageNotFoundError:
            return self._create_error_result(
                file_path,
                ParserErrorType.FILE_CORRUPTED,
                "Invalid or corrupted Word document.",
            )
        except BadZipFile:
            return self._create_error_result(
                file_path,
                ParserErrorType.FILE_CORRUPTED,
                "File is corrupted or not a valid .docx file.",
            )
        except Exception as e:
            error_msg = str(e).lower()
            if "password" in error_msg or "encrypted" in error_msg:
                return self._create_error_result(
                    file_path,
                    ParserErrorType.PASSWORD_PROTECTED,
                    "Document is password protected.",
                )
            return self._create_error_result(
                file_path,
                ParserErrorType.FILE_CORRUPTED,
                f"Failed to parse document: {e}",
            )

        return result

    def _extract_metadata(self, doc: Document, result: ParserResult) -> None:
        """Extract document metadata."""
        try:
            core_props = doc.core_properties
            result.metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception:
            result.metadata = {}

    def _extract_content(self, doc: Document, result: ParserResult) -> None:
        """Extract text content with section awareness."""
        current_section: Optional[str] = None
        section_text: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            # Check if this is a heading (section title)
            if para.style and para.style.name.startswith("Heading"):
                # Save previous section
                if section_text:
                    result.content_sections.append(
                        DocumentContent(
                            text="\n".join(section_text),
                            section_title=current_section,
                        )
                    )
                    section_text = []

                current_section = text
            else:
                section_text.append(text)

        # Save last section
        if section_text:
            result.content_sections.append(
                DocumentContent(
                    text="\n".join(section_text),
                    section_title=current_section,
                )
            )

    def _extract_tables(self, doc: Document, result: ParserResult) -> None:
        """Extract tables from document."""
        for table in doc.tables:
            parsed = self._parse_table(table)
            if parsed:
                result.tables.append(parsed)

    def _parse_table(self, table: Table) -> Optional[ParsedTable]:
        """Convert Word table to ParsedTable."""
        rows_data = []

        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            # Skip empty rows
            if any(row_cells):
                rows_data.append(row_cells)

        if len(rows_data) < 2:
            return None

        # Handle merged cells - deduplicate consecutive identical values
        cleaned_rows = []
        for row in rows_data:
            cleaned_row = []
            prev_value = None
            for value in row:
                if value != prev_value:
                    cleaned_row.append(value)
                    prev_value = value
                else:
                    # Keep placeholder for merged cell
                    cleaned_row.append("")
            cleaned_rows.append(cleaned_row)

        # Normalize row lengths
        max_cols = max(len(row) for row in cleaned_rows)
        normalized_rows = [
            row + [""] * (max_cols - len(row)) for row in cleaned_rows
        ]

        return ParsedTable(
            headers=normalized_rows[0],
            rows=normalized_rows[1:],
        )
